"""Build the Duke Coursera Specialization Design Template (.docx).

Styled to match the guidance site. Target environment is Google Docs, so this
leans on table cell shading (which survives conversion) rather than paragraph
shading (which often doesn't), and puts Heading styles inside the banner cells so
the Google Docs outline still works.
"""
import sys
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor, Emu

ICON_DIR = sys.argv[2]
OUT = sys.argv[1]

# ---- palette, lifted from the site's CSS variables ----
NAVY = "012169"
WHISPER = "F3F2F1"
HATTERAS = "E2E6ED"
GRAPHITE = "666666"
CAST_IRON = "262626"
PERSIMMON = "E89923"
WHITE = "FFFFFF"
PLACEHOLDER = "8A93A3"

TIER = {
    "spec":   {"fg": "00539B", "bg": "EFF4FB", "line": "C6D8EE", "icon": "certificate"},
    "course": {"fg": "17724A", "bg": "EDF6F1", "line": "C2DFD1", "icon": "layers"},
    "module": {"fg": "6B4E9E", "bg": "F3EFFA", "line": "D8CCEC", "icon": "list"},
    "asset":  {"fg": "B4700C", "bg": "FDF6EA", "line": "EFD9AF", "icon": "play"},
}

HEAD_FONT = "Montserrat"
BODY_FONT = "Open Sans"

CONTENT_W = 9930          # 6.9in in DXA (Letter minus 0.8in margins)


# ---------- low-level helpers ----------

def _el(tag, **attrs):
    e = OxmlElement(tag)
    for k, v in attrs.items():
        e.set(qn(f"w:{k}"), str(v))
    return e


def shade(cell, fill):
    cell._tc.get_or_add_tcPr().append(
        _el("w:shd", val="clear", color="auto", fill=fill))


def borders(cell, colour=None, sz=6, style="single", edges=("top", "left", "bottom", "right")):
    """colour=None removes borders."""
    tcb = _el("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        if edge in edges and colour:
            tcb.append(_el(f"w:{edge}", val=style, sz=sz, space=0, color=colour))
        else:
            tcb.append(_el(f"w:{edge}", val="nil"))
    cell._tc.get_or_add_tcPr().append(tcb)


def cell_margins(cell, top=90, bottom=90, left=130, right=130):
    mar = _el("w:tcMar")
    for side, val in (("top", top), ("left", left), ("bottom", bottom), ("right", right)):
        mar.append(_el(f"w:{side}", w=val, type="dxa"))
    cell._tc.get_or_add_tcPr().append(mar)


def run(p, text, *, font=BODY_FONT, size=10, bold=False, italic=False, colour=CAST_IRON):
    r = p.add_run(text)
    r.font.name = font
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    r.font.color.rgb = RGBColor.from_string(colour)
    return r


def spacing(p, before=0, after=4, line=None):
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    if line:
        pf.line_spacing = line
    return p


def mk_table(doc, rows, cols, widths):
    t = doc.add_table(rows=rows, cols=cols)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = False
    t._tbl.tblPr.append(_el("w:tblLayout", type="fixed"))
    for r in t.rows:
        for i, c in enumerate(r.cells):
            c.width = Emu(int(widths[i] * 635))     # DXA -> EMU
    return t


def blank(doc, pts=6):
    return spacing(doc.add_paragraph(), after=pts)


# ---------- components ----------

def title_block(doc, text):
    p = doc.add_paragraph()
    spacing(p, after=6)
    run(p, text, font=HEAD_FONT, size=20, bold=True, colour=NAVY)
    # navy rule under the title
    pPr = p._p.get_or_add_pPr()
    pbdr = _el("w:pBdr")
    pbdr.append(_el("w:bottom", val="single", sz=18, space=6, color=NAVY))
    pPr.append(pbdr)


def tier_strip(doc):
    labels = [("spec", "SPECIALIZATION"), ("course", "COURSE"),
              ("module", "MODULE"), ("asset", "LESSONS + ASSETS")]
    w = CONTENT_W // 4
    t = mk_table(doc, 1, 4, [w] * 4)
    for i, (tier, label) in enumerate(labels):
        c = t.cell(0, i)
        shade(c, TIER[tier]["bg"])
        borders(c, TIER[tier]["line"], sz=4)
        cell_margins(c, 70, 70, 60, 60)
        p = c.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        spacing(p, after=0)
        run(p, label, font=HEAD_FONT, size=7, bold=True, colour=TIER[tier]["fg"])
    blank(doc, 8)


def callout(doc, label, text, *, accent=NAVY, bg=WHISPER, label_colour=None):
    t = mk_table(doc, 1, 1, [CONTENT_W])
    c = t.cell(0, 0)
    shade(c, bg)
    borders(c, accent, sz=18, edges=("left",))
    cell_margins(c, 110, 110, 150, 150)
    p = c.paragraphs[0]
    spacing(p, after=2)
    if label:
        run(p, label.upper(), font=HEAD_FONT, size=7.5, bold=True,
            colour=label_colour or NAVY)
        p2 = spacing(c.add_paragraph(), after=0)
        run(p2, text, size=9.5)
    else:
        run(p, text, size=9.5)
    blank(doc, 8)


def band(doc, tier, title, sub, level):
    """Section banner: shaded cells containing a real Heading so the outline works.

    Title and sub live in separate cells so the sub text stays out of the heading —
    otherwise it shows up in the Google Docs outline alongside the title.
    """
    cfg = TIER[tier]
    t = mk_table(doc, 1, 2, [6600, CONTENT_W - 6600])
    for c in (t.cell(0, 0), t.cell(0, 1)):
        shade(c, cfg["fg"])
        borders(c, None)
    cell_margins(t.cell(0, 0), 120, 120, 150, 40)
    cell_margins(t.cell(0, 1), 120, 120, 40, 150)

    p = t.cell(0, 0).paragraphs[0]
    p.style = doc.styles[f"Heading {level}"]
    spacing(p, after=0)
    r = p.add_run()
    r.add_picture(f"{ICON_DIR}/icon-{cfg['icon']}.png", width=Pt(12))
    run(p, "  " + title, font=HEAD_FONT, size=12, bold=True, colour=WHITE)

    ps = t.cell(0, 1).paragraphs[0]
    ps.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    spacing(ps, after=0)
    if sub:
        run(ps, sub, font=BODY_FONT, size=8.5, colour="E8EEF6")
    blank(doc, 8)


def field(doc, label, placeholder, tier=None, *, hint=None, lines=1):
    fg = TIER[tier]["fg"] if tier else NAVY
    bg = TIER[tier]["bg"] if tier else WHITE
    line = TIER[tier]["line"] if tier else HATTERAS

    p = spacing(doc.add_paragraph(), before=2, after=1)
    run(p, label, font=HEAD_FONT, size=9, bold=True, colour=fg)
    if hint:
        ph = spacing(doc.add_paragraph(), after=2)
        run(ph, hint, size=8, italic=True, colour=GRAPHITE)

    t = mk_table(doc, 1, 1, [CONTENT_W])
    c = t.cell(0, 0)
    shade(c, bg)
    borders(c, line, sz=4)
    cell_margins(c, 100, 100, 140, 140)
    first = c.paragraphs[0]
    spacing(first, after=0)
    run(first, placeholder[0], size=9.5, colour=PLACEHOLDER)
    for extra in placeholder[1:]:
        pe = spacing(c.add_paragraph(), after=0)
        run(pe, extra, size=9.5, colour=PLACEHOLDER)
    for _ in range(max(0, lines - len(placeholder))):
        spacing(c.add_paragraph(), after=0)
    blank(doc, 5)


def field_pair(doc, left, right, tier=None):
    """left/right are (label, placeholder) tuples."""
    fg = TIER[tier]["fg"] if tier else NAVY
    bg = TIER[tier]["bg"] if tier else WHITE
    line = TIER[tier]["line"] if tier else HATTERAS
    half = CONTENT_W // 2

    t = mk_table(doc, 2, 2, [half, half])
    for i, (label, ph) in enumerate((left, right)):
        lc = t.cell(0, i)
        borders(lc, None)
        cell_margins(lc, 40, 30, 0, 90)
        lp = lc.paragraphs[0]
        spacing(lp, after=0)
        run(lp, label, font=HEAD_FONT, size=9, bold=True, colour=fg)

        vc = t.cell(1, i)
        shade(vc, bg)
        borders(vc, line, sz=4)
        cell_margins(vc, 100, 100, 140, 140)
        vp = vc.paragraphs[0]
        spacing(vp, after=0)
        run(vp, ph, size=9.5, colour=PLACEHOLDER)
    blank(doc, 5)


def grid(doc, headers, widths, rows, tier):
    cfg = TIER[tier]
    t = mk_table(doc, len(rows) + 1, len(headers), widths)
    for i, h in enumerate(headers):
        c = t.cell(0, i)
        shade(c, cfg["bg"])
        borders(c, cfg["line"], sz=4)
        cell_margins(c, 80, 80, 110, 110)
        p = c.paragraphs[0]
        spacing(p, after=0)
        run(p, h, font=HEAD_FONT, size=7.5, bold=True, colour=cfg["fg"])
    for ri, rowvals in enumerate(rows, start=1):
        for ci, val in enumerate(rowvals):
            c = t.cell(ri, ci)
            borders(c, cfg["line"], sz=4)
            cell_margins(c, 80, 80, 110, 110)
            p = c.paragraphs[0]
            spacing(p, after=0)
            bold = ci == 0 and val not in ("", "…")
            run(p, val, size=9,
                colour=GRAPHITE if bold else PLACEHOLDER, bold=bold)
    blank(doc, 6)


def quiz_slot(doc, course_n, module_n):
    cfg = TIER["asset"]
    t = mk_table(doc, 1, 1, [CONTENT_W])
    c = t.cell(0, 0)
    shade(c, cfg["bg"])
    borders(c, cfg["fg"], sz=12, style="dashed")
    cell_margins(c, 130, 130, 150, 150)

    p = c.paragraphs[0]
    spacing(p, after=3)
    run(p, f"Course {course_n} · Module {module_n} — Graded Quiz",
        font=HEAD_FONT, size=10, bold=True, colour="7A3300")

    p2 = spacing(c.add_paragraph(), after=6)
    run(p2, "Build this in a copy of ", size=9)
    run(p2, "Coursera Quiz Template.docx", size=9, bold=True)
    run(p2, ", then paste the link below. That file uses Coursera's import syntax "
            "exactly, so it must not be reformatted. Minimum 10 questions, with "
            "feedback for correct and incorrect responses.", size=9)

    inner = c.add_table(rows=1, cols=1)
    inner.autofit = False
    ic = inner.cell(0, 0)
    ic.width = Emu(int((CONTENT_W - 320) * 635))
    shade(ic, WHITE)
    borders(ic, cfg["line"], sz=4)
    cell_margins(ic, 90, 90, 130, 130)
    ip = ic.paragraphs[0]
    spacing(ip, after=0)
    run(ip, "[ link to completed quiz doc ]", size=9.5, colour=PLACEHOLDER)
    blank(doc, 10)


# ---------- document ----------

doc = Document()

s = doc.sections[0]
s.page_width, s.page_height = Inches(8.5), Inches(11)
for attr in ("top_margin", "bottom_margin"):
    setattr(s, attr, Inches(0.8))
s.left_margin = s.right_margin = Inches(0.8)

# Default body font
normal = doc.styles["Normal"]
normal.font.name = BODY_FONT
normal.font.size = Pt(10)
normal.font.color.rgb = RGBColor.from_string(CAST_IRON)

title_block(doc, "Coursera Specialization Design Template")
tier_strip(doc)

callout(doc, "How to use this",
        "Complete the fields and paste links to any source material created outside this "
        "document. This template is set up for three courses with three modules each — copy "
        "a whole course section or a single module section if you need more. Use the document "
        "outline to move between levels.")

p = spacing(doc.add_paragraph(), after=2)
run(p, "Prepared for:  ", font=HEAD_FONT, size=9.5, bold=True, colour=NAVY)
run(p, "[ faculty / program name ]", size=9.5, colour=PLACEHOLDER)
p = spacing(doc.add_paragraph(), after=8)
run(p, "Learning Experience Designer:  ", font=HEAD_FONT, size=9.5, bold=True, colour=NAVY)
run(p, "[ name ]", size=9.5, colour=PLACEHOLDER)

callout(doc, "Requirement",
        "Every module needs at least one graded quiz with a minimum of 10 questions and "
        "learner feedback for both correct and incorrect responses.",
        accent=PERSIMMON, bg="FDF6F0", label_colour="7A3300")

# ---- Specialization ----
band(doc, "spec", "Specialization Landing Page", "Complete once", 1)
field(doc, "Specialization title", ["[ title ]"], "spec")
field(doc, "Specialization description", ["[ description ]"], "spec",
      hint="Learner-facing. What it covers, who it's for, why it matters.", lines=3)
field_pair(doc, ("Who this is for", "[ audience ]"),
           ("Recommended background", "[ what learners should already know ]"), "spec")
field(doc, "What learners will be able to do by the end",
      ["[ overall Specialization outcome ]"], "spec")
field(doc, "Skills learners will gain",
      ["[ skill 1, skill 2, skill 3, skill 4, skill 5 ]"], "spec")
field_pair(doc, ("Estimated total learning time", "[ 12–48 hours ]"),
           ("Difficulty level", "[ Beginner / Intermediate / Advanced ]"), "spec")
field(doc, "Tools, technology, or platforms",
      ["[ any required external tools or access ]"], "spec")

p = spacing(doc.add_paragraph(), before=6, after=3)
run(p, "Courses in this Specialization", font=HEAD_FONT, size=9, bold=True,
    colour=TIER["spec"]["fg"])
grid(doc, ["#", "Course title", "Main skill / outcome", "Est. time"],
     [520, 3200, 4500, 1710],
     [[str(i), "[ title ]", "[ what learners can do ]", "[ hours ]"] for i in (1, 2, 3)],
     "spec")

field(doc, "Specialization project or culminating experience (optional)",
      ["[ description and source link, if applicable ]"], "spec")

# ---- Courses and modules ----
for cn in (1, 2, 3):
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    band(doc, "course", f"Course {cn}", "Landing page, modules, and assets", 1)
    field(doc, "Course title", ["[ title ]"], "course",
          hint="Name the skill, not a number.")
    field(doc, "Course description", ["[ learner-facing description ]"], "course", lines=3)
    field_pair(doc, ("Estimated time to complete", "[ 2–8 hours ]"),
               ("Difficulty level", "[ Beginner / Intermediate / Advanced ]"), "course")
    field_pair(doc, ("Primary domain / sub-domain", "[ domain / sub-domain ]"),
               ("Secondary domain / sub-domain", "[ domain / sub-domain ]"), "course")
    field(doc, "Tagged skills", ["[ five course-specific skills ]"], "course")
    field(doc, "Recommended background",
          ["[ what learners should know before starting ]"], "course")
    field(doc, "What learners will learn",
          ["1. [ objective ]", "2. [ objective ]",
           "3. [ objective, optional ]", "4. [ objective, optional ]"], "course",
          hint="2–4 course-level objectives.")
    field(doc, "Instructor information and bio source link",
          ["[ bio text, or link to build-ready bio and headshot ]"], "course")

    for mn in (1, 2, 3):
        band(doc, "module", f"Course {cn} · Module {mn}",
             "Duplicate this section per module", 2)
        field(doc, "Module title", ["[ title ]"], "module",
              hint="Name the skill it teaches — Coursera adds the number.")
        field(doc, "Module description", ["[ learner-facing introduction ]"], "module",
              hint="Learners see this as a message from you. Write conversationally.", lines=3)
        field(doc, "Module learning objectives",
              ["1. [ objective ]", "2. [ objective, optional ]",
               "3. [ objective, optional ]"], "module", hint="1–3 objectives.")
        field_pair(doc, ("Estimated learning time", "[ ~30+ minutes ]"),
                   ("Graded assessment title", "[ quiz title ]"), "module")

        band(doc, "asset", "Lessons & Learning Assets",
             "3–6 videos · 1–2 readings · 1–2 activities · 1 quiz", 3)
        callout(doc, None,
                "For anything developed outside this template, paste a link to the final or "
                "draft source. Your Learning Experience Designer should be able to build "
                "directly from that link.")
        grid(doc, ["#", "Lesson (optional)", "Asset type", "Asset title",
                   "Build-ready source link"],
             [520, 1900, 1700, 3200, 2610],
             [[str(i), "[ lesson ]", "[ type ]", "[ title ]", "[ link ]"]
              for i in range(1, 7)] + [["…", "", "", "", ""]],
             "asset")
        field(doc, "Additional build notes, dependencies, or accessibility information",
              ["[ notes for the learning designer ]"], lines=2)
        quiz_slot(doc, cn, mn)

p = spacing(doc.add_paragraph(), before=10, after=0)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run(p, "Copy a module section to add another module · "
       "Copy a whole course section to add another course",
    size=8.5, italic=True, colour=GRAPHITE)

doc.save(OUT)
print("wrote", OUT)
